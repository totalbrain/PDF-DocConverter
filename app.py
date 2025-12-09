import streamlit as st
import os
import io
import json
import time
import re
import base64
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional
from collections import deque

from pdf2image import convert_from_bytes
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

from models import (
    init_db, create_job, update_job_progress, complete_job, 
    fail_job, cancel_job, get_all_jobs, delete_job
)

st.set_page_config(
    page_title="PDF to DOCX Converter",
    page_icon="📄",
    layout="wide"
)

PROGRESS_FILE = "progress.json"
OUTPUT_DIR = "output"
TEMP_DIR = "temp_images"

DEFAULT_OCR_PROMPT = """You are an expert OCR + text corrector for Persian & English.
Extract and correct ALL text perfectly (spelling, grammar, punctuation).
Preserve formatting: titles with #, bold with **, lists, tables.
Output only clean corrected Markdown. No explanation."""

init_db()

def get_gemini_client():
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return None
    return genai.Client(api_key=api_key)

def save_progress(data: dict):
    with open(PROGRESS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_progress() -> Optional[dict]:
    if os.path.exists(PROGRESS_FILE):
        try:
            with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return None
    return None

def clear_progress():
    if os.path.exists(PROGRESS_FILE):
        os.remove(PROGRESS_FILE)

def image_to_bytes(image: Image.Image) -> bytes:
    buffer = io.BytesIO()
    image.save(buffer, format='PNG')
    return buffer.getvalue()

def image_to_base64(image: Image.Image) -> str:
    buffer = io.BytesIO()
    image.save(buffer, format='PNG')
    return base64.b64encode(buffer.getvalue()).decode()

def process_page_with_gemini(client, image: Image.Image, page_num: int, log_container, custom_prompt: str = None) -> tuple[str, bool]:
    image_bytes = image_to_bytes(image)
    prompt = custom_prompt if custom_prompt else DEFAULT_OCR_PROMPT
    
    max_retries = 5
    base_delay = 6
    
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    types.Part.from_bytes(
                        data=image_bytes,
                        mime_type="image/png",
                    ),
                    prompt,
                ],
            )
            
            if response.text:
                return response.text, True
            else:
                return "", True
                
        except Exception as e:
            error_str = str(e).lower()
            
            if "429" in str(e) or "resource exhausted" in error_str or "quota" in error_str:
                delay = base_delay * (2 ** attempt)
                log_container.warning(f"⚠️ Rate limit hit on page {page_num}. Waiting {delay}s... (Attempt {attempt + 1}/{max_retries})")
                time.sleep(delay)
            elif "500" in str(e) or "503" in str(e) or "internal" in error_str or "unavailable" in error_str:
                delay = base_delay * (2 ** attempt)
                log_container.warning(f"⚠️ Server error on page {page_num}. Retrying in {delay}s... (Attempt {attempt + 1}/{max_retries})")
                time.sleep(delay)
            else:
                log_container.error(f"❌ Error on page {page_num}: {str(e)[:100]}")
                return "", False
    
    log_container.error(f"❌ Failed page {page_num} after {max_retries} attempts")
    return "", False

def markdown_to_docx(markdown_pages: list[str], output_path: str):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    for page_idx, md_content in enumerate(markdown_pages):
        if page_idx > 0:
            doc.add_page_break()
        
        lines = md_content.split('\n')
        
        for line in lines:
            line = line.rstrip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                p = doc.add_paragraph(text, style='List Number')
            elif line.startswith('|') and line.endswith('|'):
                p = doc.add_paragraph(line)
            else:
                p = doc.add_paragraph()
                
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
    
    doc.save(output_path)

def markdown_to_text(markdown_pages: list[str]) -> str:
    full_text = []
    for page_idx, md_content in enumerate(markdown_pages):
        text = md_content
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'^[-*]\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)
        
        if page_idx > 0:
            full_text.append("\n\n--- Page Break ---\n\n")
        full_text.append(text)
    
    return ''.join(full_text)

def markdown_to_html(markdown_pages: list[str]) -> str:
    html_parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"><style>body{font-family:Arial,sans-serif;max-width:800px;margin:0 auto;padding:20px;}h1,h2,h3,h4{color:#333;}table{border-collapse:collapse;width:100%;}td,th{border:1px solid #ddd;padding:8px;}.page-break{page-break-after:always;border-bottom:2px dashed #ccc;margin:30px 0;}</style></head><body>']
    
    for page_idx, md_content in enumerate(markdown_pages):
        if page_idx > 0:
            html_parts.append('<div class="page-break"></div>')
        
        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                html_parts.append('<p></p>')
                continue
            
            if line.startswith('#### '):
                html_parts.append(f'<h4>{line[5:]}</h4>')
            elif line.startswith('### '):
                html_parts.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('## '):
                html_parts.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('# '):
                html_parts.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('- ') or line.startswith('* '):
                html_parts.append(f'<li>{line[2:]}</li>')
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                html_parts.append(f'<li>{text}</li>')
            else:
                line = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', line)
                line = re.sub(r'\*(.*?)\*', r'<em>\1</em>', line)
                html_parts.append(f'<p>{line}</p>')
    
    html_parts.append('</body></html>')
    return ''.join(html_parts)

def format_time(seconds: float) -> str:
    if seconds < 60:
        return f"{int(seconds)}s"
    elif seconds < 3600:
        mins = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{mins}m {secs}s"
    else:
        hours = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        return f"{hours}h {mins}m"

def format_datetime(dt):
    if dt:
        return dt.strftime("%Y-%m-%d %H:%M")
    return "-"

def show_converter_page():
    st.title("📄 PDF to DOCX Converter")
    st.markdown("Upload PDFs to extract and convert text using AI-powered OCR")
    
    client = get_gemini_client()
    if not client:
        st.error("⚠️ GEMINI_API_KEY not found in environment secrets. Please add it to continue.")
        st.info("Go to Secrets tab and add your Gemini API key with the name 'GEMINI_API_KEY'")
        return
    
    # Check database connection
    if not os.environ.get("DATABASE_URL"):
        st.error("⚠️ DATABASE_URL not found in environment secrets. Please add it to continue.")
        st.info("Go to Secrets tab and add your PostgreSQL database URL with the name 'DATABASE_URL'")
        return
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(TEMP_DIR, exist_ok=True)
    
    if 'processing' not in st.session_state:
        st.session_state.processing = False
    if 'cancel_requested' not in st.session_state:
        st.session_state.cancel_requested = False
    if 'completed_files' not in st.session_state:
        st.session_state.completed_files = []
    if 'comparison_data' not in st.session_state:
        st.session_state.comparison_data = None
    
    with st.expander("⚙️ Custom OCR Prompt", expanded=False):
        custom_prompt = st.text_area(
            "Modify the OCR instructions (leave empty for default)",
            value="",
            height=150,
            placeholder=DEFAULT_OCR_PROMPT,
            help="Customize how the AI extracts and formats text from your PDF pages"
        )
        if st.button("Reset to Default"):
            custom_prompt = ""
            st.rerun()
    
    use_prompt = custom_prompt.strip() if custom_prompt.strip() else DEFAULT_OCR_PROMPT
    
    saved_progress = load_progress()
    
    uploaded_files = st.file_uploader(
        "Choose PDF file(s)",
        type=['pdf'],
        accept_multiple_files=True,
        help="Upload one or more PDF files (up to 400-500 pages each)"
    )
    
    if uploaded_files:
        st.info(f"📁 {len(uploaded_files)} file(s) selected for processing")
    
    with st.expander("📊 Output Format Options", expanded=False):
        output_format = st.multiselect(
            "Select output formats",
            ["DOCX", "TXT", "HTML"],
            default=["DOCX"],
            help="Choose which formats to generate for each PDF"
        )
    
    if not output_format:
        output_format = ["DOCX"]
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        start_button = st.button("🚀 Start Conversion", disabled=st.session_state.processing, use_container_width=True)
    
    with col2:
        resume_enabled = saved_progress is not None and not st.session_state.processing
        resume_button = st.button("⏩ Resume", disabled=not resume_enabled, use_container_width=True)
    
    with col3:
        if st.button("🛑 Cancel", disabled=not st.session_state.processing, use_container_width=True):
            st.session_state.cancel_requested = True
            st.warning("Cancellation requested... Finishing current page...")
    
    if saved_progress and not st.session_state.processing:
        st.info(f"📝 Previous progress found: {saved_progress.get('completed_pages', 0)}/{saved_progress.get('total_pages', 0)} pages completed")
    
    if st.session_state.completed_files:
        st.subheader("📥 Download Completed Files")
        for file_info in st.session_state.completed_files:
            if os.path.exists(file_info['path']):
                with open(file_info['path'], 'rb') as f:
                    st.download_button(
                        label=f"📥 {file_info['name']}",
                        data=f.read(),
                        file_name=file_info['name'],
                        mime=file_info['mime'],
                        key=f"download_{file_info['name']}"
                    )
    
    if (start_button or resume_button) and not st.session_state.processing:
        st.session_state.processing = True
        st.session_state.cancel_requested = False
        st.session_state.completed_files = []
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        eta_text = st.empty()
        log_container = st.container()
        comparison_container = st.container()
        
        start_time = time.time()
        job_id = None
        
        try:
            files_to_process = []
            
            if resume_button and saved_progress:
                if not uploaded_files:
                    st.error("Please upload the same PDF file to resume")
                    st.session_state.processing = False
                    st.rerun()
                files_to_process = [(uploaded_files[0], saved_progress)]
            else:
                if not uploaded_files:
                    st.error("Please upload PDF file(s) first")
                    st.session_state.processing = False
                    st.rerun()
                files_to_process = [(f, None) for f in uploaded_files]
            
            total_files = len(files_to_process)
            
            for file_idx, (uploaded_file, resume_data) in enumerate(files_to_process):
                if st.session_state.cancel_requested:
                    break
                
                with log_container:
                    st.markdown(f"### Processing file {file_idx + 1}/{total_files}: {uploaded_file.name}")
                
                pdf_bytes = uploaded_file.read()
                uploaded_file.seek(0)
                
                if resume_data:
                    markdown_pages = resume_data.get('markdown_pages', [])
                    failed_pages = resume_data.get('failed_pages', [])
                    start_page = resume_data.get('completed_pages', 0)
                    total_pages = resume_data.get('total_pages', 0)
                    job_id = resume_data.get('job_id')
                    
                    with log_container:
                        st.info(f"📂 Resuming from page {start_page + 1}")
                else:
                    with log_container:
                        st.info("📄 Converting PDF pages to images...")
                    
                    images = convert_from_bytes(pdf_bytes, dpi=300)
                    total_pages = len(images)
                    
                    with log_container:
                        st.success(f"✅ Found {total_pages} pages")
                    
                    markdown_pages = [""] * total_pages
                    failed_pages = []
                    start_page = 0
                    
                    with log_container:
                        st.info(f"🔍 Creating job in database for {uploaded_file.name}...")
                        st.info(f"DATABASE_URL exists: {bool(os.environ.get('DATABASE_URL'))}")
                    
                    job_id = create_job(uploaded_file.name, total_pages, use_prompt if use_prompt != DEFAULT_OCR_PROMPT else None)
                    
                    with log_container:
                        st.info(f"Job ID returned: {job_id}")
                    
                    if not job_id:
                        with log_container:
                            st.error("❌ Failed to create job in database.")
                            st.error("Please check:")
                            st.error("1. DATABASE_URL is set in Secrets")
                            st.error("2. Database is accessible")
                            st.error("3. Check console logs for more details")
                        st.session_state.processing = False
                        st.rerun()
                    
                    save_progress({
                        'total_pages': total_pages,
                        'completed_pages': 0,
                        'markdown_pages': markdown_pages,
                        'failed_pages': failed_pages,
                        'filename': uploaded_file.name,
                        'job_id': job_id
                    })
                
                images = convert_from_bytes(pdf_bytes, dpi=300)
                
                page_times = deque(maxlen=10)
                last_image = None
                last_text = None
                
                for page_idx in range(start_page, total_pages):
                    if st.session_state.cancel_requested:
                        if job_id:
                            cancel_job(job_id, page_idx)
                        with log_container:
                            st.warning(f"⚠️ Cancelled at page {page_idx + 1}")
                        break
                    
                    page_start_time = time.time()
                    
                    overall_progress = (file_idx * total_pages + page_idx + 1) / (total_files * total_pages)
                    progress_bar.progress(overall_progress)
                    status_text.text(f"File {file_idx + 1}/{total_files} | Page {page_idx + 1}/{total_pages} ({int(overall_progress * 100)}%)")
                    
                    if page_times:
                        avg_time = sum(page_times) / len(page_times)
                        remaining_pages = total_pages - page_idx - 1
                        for remaining_file_idx in range(file_idx + 1, total_files):
                            remaining_pages += total_pages
                        eta_seconds = remaining_pages * avg_time
                        eta_text.text(f"⏱️ ETA: {format_time(eta_seconds)}")
                    
                    image = images[page_idx]
                    
                    text, success = process_page_with_gemini(client, image, page_idx + 1, log_container, use_prompt)
                    
                    if success:
                        markdown_pages[page_idx] = text
                        last_image = image
                        last_text = text
                        with log_container:
                            st.success(f"✅ Page {page_idx + 1} processed")
                    else:
                        failed_pages.append(page_idx)
                        with log_container:
                            st.warning(f"⚠️ Page {page_idx + 1} queued for retry")
                    
                    if (page_idx + 1) % 10 == 0:
                        save_progress({
                            'total_pages': total_pages,
                            'completed_pages': page_idx + 1,
                            'markdown_pages': markdown_pages,
                            'failed_pages': failed_pages,
                            'filename': uploaded_file.name,
                            'job_id': job_id
                        })
                        if job_id:
                            update_job_progress(job_id, page_idx + 1, len(failed_pages))
                    
                    page_time = time.time() - page_start_time
                    page_times.append(page_time)
                    
                    if page_idx < total_pages - 1:
                        time.sleep(6)
                
                if not st.session_state.cancel_requested and failed_pages:
                    with log_container:
                        st.info(f"🔄 Retrying {len(failed_pages)} failed pages...")
                    
                    retry_failed = []
                    for fail_idx in failed_pages:
                        if st.session_state.cancel_requested:
                            break
                        
                        image = images[fail_idx]
                        text, success = process_page_with_gemini(client, image, fail_idx + 1, log_container, use_prompt)
                        
                        if success:
                            markdown_pages[fail_idx] = text
                            with log_container:
                                st.success(f"✅ Page {fail_idx + 1} recovered")
                        else:
                            retry_failed.append(fail_idx)
                        
                        time.sleep(6)
                    
                    if retry_failed:
                        with log_container:
                            st.warning(f"⚠️ {len(retry_failed)} pages could not be recovered: {[p+1 for p in retry_failed]}")
                    
                    failed_pages = retry_failed
                
                if not st.session_state.cancel_requested:
                    with log_container:
                        st.info("📝 Creating output documents...")
                    
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    base_name = Path(uploaded_file.name).stem
                    
                    processing_time = time.time() - start_time
                    
                    output_files = []
                    
                    if "DOCX" in output_format:
                        output_filename = f"{base_name}_{timestamp}.docx"
                        output_path = os.path.join(OUTPUT_DIR, output_filename)
                        markdown_to_docx(markdown_pages, output_path)
                        output_files.append({
                            'path': output_path,
                            'name': output_filename,
                            'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        })
                        main_output = output_path
                    
                    if "TXT" in output_format:
                        txt_filename = f"{base_name}_{timestamp}.txt"
                        txt_path = os.path.join(OUTPUT_DIR, txt_filename)
                        txt_content = markdown_to_text(markdown_pages)
                        with open(txt_path, 'w', encoding='utf-8') as f:
                            f.write(txt_content)
                        output_files.append({
                            'path': txt_path,
                            'name': txt_filename,
                            'mime': 'text/plain'
                        })
                        if "DOCX" not in output_format:
                            main_output = txt_path
                    
                    if "HTML" in output_format:
                        html_filename = f"{base_name}_{timestamp}.html"
                        html_path = os.path.join(OUTPUT_DIR, html_filename)
                        html_content = markdown_to_html(markdown_pages)
                        with open(html_path, 'w', encoding='utf-8') as f:
                            f.write(html_content)
                        output_files.append({
                            'path': html_path,
                            'name': html_filename,
                            'mime': 'text/html'
                        })
                        if "DOCX" not in output_format and "TXT" not in output_format:
                            main_output = html_path
                    
                    st.session_state.completed_files.extend(output_files)
                    
                    if job_id:
                        complete_job(job_id, main_output, processing_time, len(failed_pages), output_files)
                    
                    clear_progress()
                    
                    with log_container:
                        st.success(f"🎉 {uploaded_file.name} converted successfully!")
                    
                    if last_image and last_text:
                        st.session_state.comparison_data = {
                            'image': last_image,
                            'text': last_text,
                            'page': total_pages
                        }
            
            progress_bar.progress(1.0)
            status_text.text("✅ All conversions complete!")
            eta_text.empty()
        
        except Exception as e:
            with log_container:
                st.error(f"❌ Error: {str(e)}")
            
            if job_id:
                fail_job(job_id, str(e))
            
            if "quota" in str(e).lower() or "exceeded" in str(e).lower():
                st.error("⚠️ API quota exceeded. Please wait or upgrade your Gemini API plan.")
        
        finally:
            st.session_state.processing = False
            st.rerun()
    
    if st.session_state.comparison_data:
        st.subheader("🔍 OCR Quality Comparison (Last Page)")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Original Page Image**")
            img = st.session_state.comparison_data['image']
            thumb = img.copy()
            thumb.thumbnail((400, 600))
            st.image(thumb, use_container_width=True)
        
        with col2:
            st.markdown("**Extracted Text**")
            st.text_area(
                "OCR Result",
                value=st.session_state.comparison_data['text'],
                height=400,
                disabled=True,
                label_visibility="collapsed"
            )

def show_history_page():
    st.title("📋 Conversion History")
    st.markdown("View and manage your past conversion jobs")
    
    jobs = get_all_jobs()
    
    if not jobs:
        st.info("No conversion jobs found. Start converting PDFs to see them here!")
        return
    
    st.metric("Total Jobs", len(jobs))
    
    for job in jobs:
        status_emoji = {
            'completed': '✅',
            'processing': '🔄',
            'failed': '❌',
            'cancelled': '⚠️',
            'pending': '⏳'
        }.get(job['status'], '❓')
        
        with st.expander(f"{status_emoji} {job['filename']} - {format_datetime(job['created_at'])}"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.write(f"**Status:** {job['status'].title()}")
                st.write(f"**Pages:** {job['completed_pages']}/{job['total_pages']}")
            
            with col2:
                st.write(f"**Created:** {format_datetime(job['created_at'])}")
                st.write(f"**Completed:** {format_datetime(job['completed_at'])}")
            
            with col3:
                if job['processing_time_seconds']:
                    st.write(f"**Duration:** {format_time(job['processing_time_seconds'])}")
                if job['failed_pages']:
                    st.write(f"**Failed Pages:** {job['failed_pages']}")
            
            if job['custom_prompt']:
                st.text_area("Custom Prompt Used", value=job['custom_prompt'], height=100, disabled=True)
            
            if job['error_message']:
                st.error(f"Error: {job['error_message']}")
            
            if job['status'] == 'completed':
                output_paths = job.get('output_paths', [])
                if output_paths:
                    st.markdown("**Download Files:**")
                    download_cols = st.columns(min(len(output_paths), 3))
                    for idx, file_info in enumerate(output_paths):
                        file_path = file_info.get('path', '')
                        if file_path and os.path.exists(file_path):
                            with download_cols[idx % 3]:
                                with open(file_path, 'rb') as f:
                                    st.download_button(
                                        label=f"📥 {file_info.get('name', 'Download')}",
                                        data=f.read(),
                                        file_name=file_info.get('name', os.path.basename(file_path)),
                                        mime=file_info.get('mime', 'application/octet-stream'),
                                        key=f"redownload_{job['id']}_{idx}"
                                    )
                elif job['output_path'] and os.path.exists(job['output_path']):
                    with open(job['output_path'], 'rb') as f:
                        st.download_button(
                            label="📥 Re-download",
                            data=f.read(),
                            file_name=os.path.basename(job['output_path']),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"redownload_{job['id']}"
                        )
            
            if st.button("🗑️ Delete", key=f"delete_{job['id']}"):
                delete_job(job['id'])
                st.rerun()

def main():
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["🔄 Converter", "📋 History"])
    
    if page == "🔄 Converter":
        show_converter_page()
    else:
        show_history_page()

if __name__ == "__main__":
    main()
